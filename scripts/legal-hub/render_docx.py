#!/usr/bin/env python3
"""Render markdown draft into court-style DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install with: pip install python-docx"
    ) from exc

# Regex for footnote reference [^1] and definition [^1]: text
_FOOTNOTE_REF_RE = re.compile(r"\[\^(\w+)\]")
_FOOTNOTE_DEF_RE = re.compile(r"^\[\^(\w+)\]:\s+(.+)$")
# Regex for table row: | col | col | (must start and end with |)
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
# Regex for separator row: |---|---| or |:---:|:---|
_TABLE_SEP_RE = re.compile(r"^\|[\s:]*-{3,}[\s:]*(\|[\s:]*-{3,}[\s:]*)*\|$")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Render a markdown draft to DOCX.")
    parser.add_argument("matter_path", help="Matter folder path")
    parser.add_argument("--input", default="03_drafts/draft.md", help="Markdown input path (relative to matter)")
    parser.add_argument("--output", default="04_final/final.docx", help="DOCX output path (relative to matter)")
    parser.add_argument("--font", default="Malgun Gothic", help="Main document font")
    parser.add_argument("--font-size", type=float, default=12.0, help="Main font size (pt)")
    parser.add_argument("--line-spacing", type=float, default=1.6, help="Line spacing multiplier")
    parser.add_argument("--top-cm", type=float, default=4.5, help="Top margin in cm")
    parser.add_argument("--bottom-cm", type=float, default=3.0, help="Bottom margin in cm")
    parser.add_argument("--left-cm", type=float, default=2.0, help="Left margin in cm")
    parser.add_argument("--right-cm", type=float, default=2.0, help="Right margin in cm")
    return parser.parse_args()


def configure_document(doc: Document, font_name: str, font_size: float, line_spacing: float, margins: dict[str, float]) -> None:
    for section in doc.sections:
        section.top_margin = Cm(margins["top"])
        section.bottom_margin = Cm(margins["bottom"])
        section.left_margin = Cm(margins["left"])
        section.right_margin = Cm(margins["right"])

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)


def add_heading(doc: Document, text: str, level: int) -> None:
    size_map = {1: 16, 2: 14, 3: 13}
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text.strip())
    run.bold = True
    run.font.size = Pt(size_map.get(level, 12))
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_runs_with_bold(paragraph, text: str) -> None:
    """Parse **bold** markers and add runs with appropriate formatting."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            if part:
                paragraph.add_run(part)


def add_plain_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs_with_bold(paragraph, text.strip())


def add_list_item(doc: Document, text: str, numbered: bool) -> None:
    style_name = "List Number" if numbered else "List Bullet"
    paragraph = doc.add_paragraph(style=style_name)
    add_runs_with_bold(paragraph, text.strip())


# ── Table support ────────────────────────────────────────────


def _parse_table_cells(row_text: str) -> list[str]:
    """Split a pipe-delimited row into cell texts, stripping whitespace."""
    # Remove leading/trailing pipes, then split by |
    inner = row_text.strip().strip("|")
    return [cell.strip() for cell in inner.split("|")]


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a styled table to the document."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        para = table.cell(0, i).paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.bold = True

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx in range(n_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            para = table.cell(row_idx + 1, col_idx).paragraphs[0]
            para.clear()
            add_runs_with_bold(para, cell_text)


# ── Footnote support ─────────────────────────────────────────


def collect_footnote_defs(lines: list[str]) -> dict[str, str]:
    """Scan lines for [^key]: definition and return {key: text} map."""
    defs: dict[str, str] = {}
    for line in lines:
        m = _FOOTNOTE_DEF_RE.match(line.strip())
        if m:
            defs[m.group(1)] = m.group(2)
    return defs


def add_footnote_paragraph(doc: Document, text: str, fn_defs: dict[str, str]) -> None:
    """Add a paragraph with [^N] replaced by superscript numbers."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Split text around footnote refs
    parts = _FOOTNOTE_REF_RE.split(text)
    # parts alternates: text, key, text, key, ...
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Normal text segment — handle **bold**
            if part:
                add_runs_with_bold(paragraph, part)
        else:
            # Footnote key — render as superscript
            run = paragraph.add_run(f"{part})")
            run.font.superscript = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x99)


def _is_table_row(line: str) -> bool:
    """Check if a line looks like a markdown table row (not code with pipes)."""
    stripped = line.strip()
    if not _TABLE_ROW_RE.match(stripped):
        return False
    # Exclude lines that are inside backtick code spans
    # Simple heuristic: if backticks wrap the pipes, it's code
    if stripped.startswith("`") or stripped.count("`") >= 2:
        no_code = re.sub(r"`[^`]+`", "", stripped)
        return bool(_TABLE_ROW_RE.match(no_code.strip()))
    return True


def render_markdown(doc: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    fn_defs = collect_footnote_defs(lines)
    has_footnotes = bool(fn_defs)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # Skip footnote definition lines (rendered at end)
        if _FOOTNOTE_DEF_RE.match(stripped):
            i += 1
            continue

        # Empty line
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # ── Table block ──
        if _is_table_row(stripped):
            # Collect consecutive table lines
            table_lines: list[str] = []
            while i < len(lines) and _is_table_row(lines[i].rstrip()):
                table_lines.append(lines[i].rstrip().strip())
                i += 1
            # Also grab separator if next
            if i < len(lines) and _TABLE_SEP_RE.match(lines[i].strip()):
                table_lines.append(lines[i].rstrip().strip())
                i += 1
                while i < len(lines) and _is_table_row(lines[i].rstrip()):
                    table_lines.append(lines[i].rstrip().strip())
                    i += 1

            # Parse: first line = headers, skip separator, rest = data
            headers = _parse_table_cells(table_lines[0])
            rows: list[list[str]] = []
            for tl in table_lines[1:]:
                if _TABLE_SEP_RE.match(tl):
                    continue
                rows.append(_parse_table_cells(tl))
            add_table(doc, headers, rows)
            continue

        # ── Heading ──
        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            add_heading(doc, heading_match.group(2), level)
            i += 1
            continue

        # ── Bullet list ──
        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            add_list_item(doc, bullet_match.group(1), numbered=False)
            i += 1
            continue

        # ── Numbered list ──
        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            add_list_item(doc, number_match.group(1), numbered=True)
            i += 1
            continue

        # ── Plain paragraph (with footnote support) ──
        if has_footnotes and _FOOTNOTE_REF_RE.search(stripped):
            add_footnote_paragraph(doc, stripped, fn_defs)
        else:
            add_plain_paragraph(doc, stripped)
        i += 1

    # ── Append footnote section at end ──
    if fn_defs:
        doc.add_paragraph()
        add_heading(doc, "주석", 3)
        for key, text in fn_defs.items():
            para = doc.add_paragraph()
            ref_run = para.add_run(f"{key}) ")
            ref_run.font.superscript = True
            ref_run.font.size = Pt(9)
            add_runs_with_bold(para, text)


def main() -> int:
    args = parse_args()
    matter_root = Path(args.matter_path)
    input_path = matter_root / args.input
    output_path = matter_root / args.output

    if not input_path.exists():
        raise SystemExit(f"Input markdown not found: {input_path}")

    markdown_text = input_path.read_text(encoding="utf-8")
    doc = Document()
    configure_document(
        doc=doc,
        font_name=args.font,
        font_size=args.font_size,
        line_spacing=args.line_spacing,
        margins={
            "top": args.top_cm,
            "bottom": args.bottom_cm,
            "left": args.left_cm,
            "right": args.right_cm,
        },
    )
    render_markdown(doc, markdown_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"DOCX generated: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
