#!/usr/bin/env python3
"""Tests for render_docx.py — table and footnote rendering."""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from render_docx import (
    add_table,
    add_footnote_paragraph,
    collect_footnote_defs,
    render_markdown,
)


# ── Table Tests ──────────────────────────────────────────────


class TestAddTable:
    def test_basic_table(self):
        doc = Document()
        headers = ["항목", "결과", "상태"]
        rows = [
            ["Claude Code 버전", "v2.1.47", "PASS"],
            ["Chrome Extension", "등록됨", "PASS"],
        ]
        add_table(doc, headers, rows)

        table = doc.tables[0]
        assert len(table.rows) == 3  # 1 header + 2 data
        assert len(table.columns) == 3
        assert table.cell(0, 0).text == "항목"
        assert table.cell(1, 0).text == "Claude Code 버전"
        assert table.cell(2, 2).text == "PASS"

    def test_header_row_is_bold(self):
        doc = Document()
        add_table(doc, ["A", "B"], [["1", "2"]])
        header_cell = doc.tables[0].cell(0, 0)
        run = header_cell.paragraphs[0].runs[0]
        assert run.bold is True

    def test_single_column_table(self):
        doc = Document()
        add_table(doc, ["항목"], [["값1"], ["값2"]])
        table = doc.tables[0]
        assert len(table.columns) == 1
        assert len(table.rows) == 3

    def test_bold_in_cell(self):
        doc = Document()
        add_table(doc, ["col"], [["**강조**"]])
        cell = doc.tables[0].cell(1, 0)
        runs = cell.paragraphs[0].runs
        assert any(r.bold for r in runs)

    def test_empty_cell(self):
        doc = Document()
        add_table(doc, ["A", "B"], [["값", ""]])
        assert doc.tables[0].cell(1, 1).text == ""


# ── Footnote Tests ───────────────────────────────────────────


class TestCollectFootnoteDefs:
    def test_basic_definition(self):
        lines = [
            "본문 텍스트",
            "",
            "[^1]: 대법원 2024. 1. 1. 선고 2023다12345 판결",
            "[^2]: 근로기준법 제23조 제1항",
        ]
        defs = collect_footnote_defs(lines)
        assert defs["1"] == "대법원 2024. 1. 1. 선고 2023다12345 판결"
        assert defs["2"] == "근로기준법 제23조 제1항"

    def test_no_definitions(self):
        lines = ["본문만 있는 텍스트", "두 번째 줄"]
        defs = collect_footnote_defs(lines)
        assert defs == {}


class TestAddFootnoteParagraph:
    def test_text_without_footnotes(self):
        doc = Document()
        add_footnote_paragraph(doc, "일반 텍스트입니다.", {})
        assert doc.paragraphs[0].text == "일반 텍스트입니다."

    def test_footnote_ref_becomes_superscript(self):
        doc = Document()
        defs = {"1": "참고 문헌"}
        add_footnote_paragraph(doc, "해고는 부당하다[^1].", defs)
        para = doc.paragraphs[0]
        # Should have runs: "해고는 부당하다", superscript "1)", "."
        texts = [r.text for r in para.runs]
        joined = "".join(texts)
        assert "해고는 부당하다" in joined
        # Superscript run exists
        superscript_runs = [r for r in para.runs if r.font.superscript]
        assert len(superscript_runs) >= 1

    def test_multiple_footnotes(self):
        doc = Document()
        defs = {"1": "판례A", "2": "판례B"}
        add_footnote_paragraph(doc, "주장A[^1]과 주장B[^2].", defs)
        para = doc.paragraphs[0]
        superscript_runs = [r for r in para.runs if r.font.superscript]
        assert len(superscript_runs) >= 2


# ── Integration: render_markdown with tables ─────────────────


class TestRenderMarkdownTable:
    TABLE_MD = (
        "| 플랫폼 | 결과 |\n"
        "|--------|------|\n"
        "| 엘박스 | 9,867건 |\n"
        "| 빅케이스 | 3,508건 |\n"
    )

    def test_table_parsed_in_markdown(self):
        doc = Document()
        render_markdown(doc, self.TABLE_MD)
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "플랫폼"
        assert doc.tables[0].cell(1, 0).text == "엘박스"

    def test_text_before_and_after_table(self):
        md = "제목 문단\n\n" + self.TABLE_MD + "\n후속 문단\n"
        doc = Document()
        render_markdown(doc, md)
        assert len(doc.tables) == 1
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "제목 문단" in texts
        assert "후속 문단" in texts

    def test_multiple_tables(self):
        md = self.TABLE_MD + "\n중간 텍스트\n\n" + self.TABLE_MD
        doc = Document()
        render_markdown(doc, md)
        assert len(doc.tables) == 2


class TestRenderMarkdownFootnote:
    def test_footnote_in_paragraph(self):
        md = (
            "해고가 부당하다[^1].\n"
            "\n"
            "[^1]: 대법원 2024. 1. 1. 선고 판결\n"
        )
        doc = Document()
        render_markdown(doc, md)
        # Footnote definition lines should not appear as paragraphs
        visible_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert not any("[^1]:" in t for t in visible_texts)

    def test_footnote_section_appended(self):
        md = (
            "본문[^1].\n"
            "\n"
            "[^1]: 각주 내용\n"
        )
        doc = Document()
        render_markdown(doc, md)
        # Footnote section heading should exist
        texts = [p.text for p in doc.paragraphs]
        assert any("각주" in t or "주석" in t for t in texts)


# ── Edge Cases ───────────────────────────────────────────────


class TestEdgeCases:
    def test_separator_only_row_ignored(self):
        """A separator row like |---|---| should not become a data row."""
        md = "| A | B |\n|---|---|\n| 1 | 2 |\n"
        doc = Document()
        render_markdown(doc, md)
        table = doc.tables[0]
        assert len(table.rows) == 2  # header + 1 data, no separator row

    def test_pipe_in_code_not_parsed_as_table(self):
        """Inline code with pipes should not trigger table parsing."""
        md = "명령어: `echo a | grep b`\n"
        doc = Document()
        render_markdown(doc, md)
        assert len(doc.tables) == 0

    def test_colon_alignment_separators(self):
        """Separators like |:---:|:---| should be handled."""
        md = "| A | B |\n|:---:|:---|\n| centered | left |\n"
        doc = Document()
        render_markdown(doc, md)
        assert len(doc.tables) == 1
